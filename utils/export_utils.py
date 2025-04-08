"""
Utility functions for exporting questionnaires to different formats
"""

import os
import tempfile
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, ListItem, ListFlowable
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_JUSTIFY


def export_to_word(questionnaire, research_topic, research_description, target_audience, questionnaire_purpose):
    """
    Export questionnaire to Word document.
    
    Args:
        questionnaire (list): List of sections with questions
        research_topic (str): Title of the questionnaire
        research_description (str): Description of the research
        target_audience (str): Target audience for the questionnaire
        questionnaire_purpose (str): Purpose of the questionnaire
        
    Returns:
        str: Path to the generated Word document
    """
    doc = Document()
    
    # Set up document styles
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)
    
    # Add title
    title = doc.add_heading(research_topic, level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add metadata
    doc.add_paragraph(f"Target Audience: {target_audience}")
    doc.add_paragraph(f"Purpose: {questionnaire_purpose}")
    
    # Add introduction
    doc.add_heading("Introduction", level=2)
    intro_text = f"Thank you for participating in this {questionnaire_purpose.lower()}. This questionnaire aims to {research_description}. Your responses will be kept confidential and will be used only for research purposes."
    doc.add_paragraph(intro_text)
    
    # Add sections and questions
    for section_index, section in enumerate(questionnaire):
        # Add section header
        doc.add_heading(section['title'], level=2)
        
        # Add section description if available
        if section.get('description'):
            doc.add_paragraph(section['description'])
        
        # Add questions
        for question_index, question in enumerate(section['questions']):
            q_num = question_index + 1
            
            # Add question text
            p = doc.add_paragraph()
            p.add_run(f"{q_num}. {question['text']}").bold = True
            p.add_run(f" ({question['type']})").italic = True
            
            # Add options based on question type
            if question['type'] in ['Multiple Choice', 'Checkbox'] and 'options' in question:
                for option_index, option in enumerate(question['options']):
                    doc.add_paragraph(f"   □ {option}", style='List Bullet')
            
            elif question['type'] == 'Likert Scale':
                table = doc.add_table(rows=2, cols=5)
                table.style = 'Table Grid'
                
                # Add headers
                cells = table.rows[0].cells
                cells[0].text = "Strongly Disagree"
                cells[1].text = ""
                cells[2].text = ""
                cells[3].text = ""
                cells[4].text = "Strongly Agree"
                
                # Add numbers
                cells = table.rows[1].cells
                cells[0].text = "1"
                cells[1].text = "2"
                cells[2].text = "3"
                cells[3].text = "4"
                cells[4].text = "5"
            
            elif question['type'] == 'Rating':
                table = doc.add_table(rows=2, cols=5)
                table.style = 'Table Grid'
                
                # Add headers
                cells = table.rows[0].cells
                cells[0].text = "Poor"
                cells[1].text = ""
                cells[2].text = ""
                cells[3].text = ""
                cells[4].text = "Excellent"
                
                # Add numbers
                cells = table.rows[1].cells
                cells[0].text = "1"
                cells[1].text = "2"
                cells[2].text = "3"
                cells[3].text = "4"
                cells[4].text = "5"
            
            elif question['type'] == 'Open-Ended':
                doc.add_paragraph("Answer: ________________________________")
                doc.add_paragraph("________________________________")
                doc.add_paragraph("________________________________")
            
            # Add spacing between questions
            doc.add_paragraph()
    
    # Add thank you message
    doc.add_heading("Thank You", level=2)
    doc.add_paragraph("Thank you for completing this questionnaire. Your feedback is valuable for our research.")
    
    # Save the document to a temporary file
    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
    doc.save(temp_file.name)
    temp_file.close()
    
    return temp_file.name


def export_to_pdf(questionnaire, research_topic, research_description, target_audience, questionnaire_purpose):
    """
    Export questionnaire to PDF document.
    
    Args:
        questionnaire (list): List of sections with questions
        research_topic (str): Title of the questionnaire
        research_description (str): Description of the research
        target_audience (str): Target audience for the questionnaire
        questionnaire_purpose (str): Purpose of the questionnaire
        
    Returns:
        str: Path to the generated PDF document
    """
    # Create a temporary file
    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.pdf')
    temp_file.close()
    
    # Set up the document
    doc = SimpleDocTemplate(
        temp_file.name,
        pagesize=letter,
        rightMargin=72,
        leftMargin=72,
        topMargin=72,
        bottomMargin=72
    )
    
    # Get styles
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(
        name='QuestionnaireTitle',
        parent=styles['Heading1'],
        fontSize=18,
        alignment=TA_CENTER,
        spaceAfter=12
    ))
    styles.add(ParagraphStyle(
        name='QuestionnaireSectionTitle',
        parent=styles['Heading2'],
        fontSize=14,
        spaceBefore=12,
        spaceAfter=6
    ))
    styles.add(ParagraphStyle(
        name='QuestionText',
        parent=styles['Normal'],
        fontSize=11,
        fontName='Helvetica-Bold',
        spaceBefore=8,
        spaceAfter=4
    ))
    styles.add(ParagraphStyle(
        name='OptionText',
        parent=styles['Normal'],
        fontSize=10,
        leftIndent=20
    ))
    
    # Create story
    story = []
    
    # Add title
    story.append(Paragraph(research_topic, styles['QuestionnaireTitle']))
    story.append(Spacer(1, 12))
    
    # Add metadata
    story.append(Paragraph(f"<b>Target Audience:</b> {target_audience}", styles['Normal']))
    story.append(Paragraph(f"<b>Purpose:</b> {questionnaire_purpose}", styles['Normal']))
    story.append(Spacer(1, 12))
    
    # Add introduction
    story.append(Paragraph("Introduction", styles['QuestionnaireSectionTitle']))
    intro_text = f"Thank you for participating in this {questionnaire_purpose.lower()}. This questionnaire aims to {research_description}. Your responses will be kept confidential and will be used only for research purposes."
    story.append(Paragraph(intro_text, styles['Normal']))
    story.append(Spacer(1, 12))
    
    # Add sections and questions
    for section_index, section in enumerate(questionnaire):
        # Add section header
        story.append(Paragraph(section['title'], styles['QuestionnaireSectionTitle']))
        
        # Add section description if available
        if section.get('description'):
            story.append(Paragraph(section['description'], styles['Normal']))
            story.append(Spacer(1, 6))
        
        # Add questions
        for question_index, question in enumerate(section['questions']):
            q_num = question_index + 1
            
            # Add question text
            story.append(Paragraph(
                f"{q_num}. {question['text']} <i>({question['type']})</i>", 
                styles['QuestionText']
            ))
            
            # Add options based on question type
            if question['type'] in ['Multiple Choice', 'Checkbox'] and 'options' in question:
                options_list = []
                for option in question['options']:
                    checkbox = "□ " if question['type'] == 'Multiple Choice' else "☐ "
                    options_list.append(ListItem(Paragraph(f"{checkbox}{option}", styles['OptionText'])))
                
                story.append(ListFlowable(
                    options_list,
                    bulletType='bullet',
                    start='',
                    bulletFontName='Helvetica',
                    bulletFontSize=10,
                    leftIndent=20
                ))
            
            elif question['type'] in ['Likert Scale', 'Rating']:
                # Create table for scale
                is_likert = question['type'] == 'Likert Scale'
                header1 = "Strongly Disagree" if is_likert else "Poor"
                header2 = "Strongly Agree" if is_likert else "Excellent"
                
                data = [
                    [header1, '', '', '', header2],
                    ['1', '2', '3', '4', '5']
                ]
                
                # Create the table
                table = Table(data, colWidths=[80, 80, 80, 80, 80])
                table.setStyle(TableStyle([
                    ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
                    ('BACKGROUND', (0, 0), (-1, 0), colors.lightgrey),
                    ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                    ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, -1), 9),
                ]))
                
                story.append(table)
            
            elif question['type'] == 'Open-Ended':
                # Create lines for open-ended response
                data = [[''], [''], ['']]
                table = Table(data, colWidths=[400])
                table.setStyle(TableStyle([
                    ('LINEBELOW', (0, 0), (0, 2), 0.5, colors.grey),
                    ('TOPPADDING', (0, 0), (0, 2), 15),
                ]))
                story.append(table)
            
            # Add spacing between questions
            story.append(Spacer(1, 10))
    
    # Add thank you message
    story.append(Paragraph("Thank You", styles['QuestionnaireSectionTitle']))
    story.append(Paragraph("Thank you for completing this questionnaire. Your feedback is valuable for our research.", styles['Normal']))
    
    # Build the PDF
    doc.build(story)
    
    return temp_file.name 